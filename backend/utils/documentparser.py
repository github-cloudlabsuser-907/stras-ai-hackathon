
from typing import Any
import io

import requests, PyPDF2
from PyPDF2 import PdfReader

from reportlab.platypus import SimpleDocTemplate, Paragraph as PDFParagraph, Table as PDFTable
from reportlab.lib.styles import getSampleStyleSheet

from .azureblobstorage import AzureBlobStorageClient
from docx import Document
from docx.oxml import CT_Picture
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from urllib.parse import unquote


class DocumentParser:
    """
    A class for parsing documents, extracting sections, and creating PDF streams from the document content.

    Attributes:
        file_name (str): The name of the document file.
        pdf_doc_stream: The PDF document stream.
        blob_client (AzureBlobStorageClient): An instance of AzureBlobStorageClient.
    """

    def __init__(self, file_name: str, pdf_doc_stream, blob_client: AzureBlobStorageClient = None):
        """
        Initialize the DocumentParser class.

        Args:
            file_name (str): The name of the document file.
            pdf_doc_stream: The PDF document stream.
            blob_client (AzureBlobStorageClient, optional): An instance of AzureBlobStorageClient. Defaults to None.
        """
        self.file_name = file_name 
        self.pdf_doc_stream = pdf_doc_stream
        self.blob_client: AzureBlobStorageClient = AzureBlobStorageClient() if blob_client is None else blob_client
    
    
    def iter_block_items(self, parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        """
        from docx.document import Document
        if isinstance(parent, Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("Something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def get_level(self, paragraph)-> str|None:
        """
        Function to get the level of a paragraph.
        """
        style = paragraph.style.name
        if style.startswith("Heading 4"):
            return "header_4"
        elif style.startswith("Heading 3"):
            return "header_3"
        elif style.startswith("Heading 2"):
            return "header_2"
        elif style.startswith("Heading 1"):
            return "header_1"
        else:
            return None
        
    
    def pdf_to_docx(self, pdf_bytes):

        text = ""
        doc = Document()
        with io.BytesIO(pdf_bytes) as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
                doc.add_paragraph(text)

        return doc
    
    # def create_docx_from_pdf(self, pdf_bytes):
    #     print("PDF BYTES")
    #     text = self.pdf_to_text(pdf_bytes)
        
    #     doc.add_paragraph(text)
    #     return doc
        
    def parse_document(self):
        chapter_structure = []
        doc = Document(self.pdf_doc_stream)
        # Iterate through paragraphs and store the hierarchical structure
        current_sections = [0, 0, 0, 0]  # To track the current section number at each level
        for item in self.iter_block_items(doc):
            if isinstance(item, Paragraph):
                level = self.get_level(item)
                if level:
                    level_index = int(level[-1]) - 1  # Convert 'header_#' to an index
                    for i in range(level_index + 1, 4):
                        current_sections[i] = 0  # Reset lower level sections
                    current_sections[level_index] += 1  # Increment the corresponding section number
                    section_title = '.'.join(str(num) for num in current_sections[:level_index + 1]) + " " + item.text
                    chapter_structure.append({"title": section_title, "content": [], "tables": []})
                elif chapter_structure:  # Add paragraph content to the latest section
                    chapter_structure[-1]["content"].append(item.text)

            elif isinstance(item, Table) and chapter_structure:
                    chapter_structure[-1]["tables"].append(item)

        return chapter_structure
        # # List to store the hierarchical structure of the document
        # chapter_structure = []
        # doc = self.pdf_to_docx(self.pdf_doc_stream)
        # # Iterate through paragraphs and store the hierarchical structure
        # current_sections = [0, 0, 0, 0]  # To track the current section number at each level
        # for item in self.iter_block_items(doc):
        #     if isinstance(item, Paragraph):
        #         level = self.get_level(item)
        #         if level:
        #             level_index = int(level[-1]) - 1  # Convert 'header_#' to an index
        #             for i in range(level_index + 1, 4):
        #                 current_sections[i] = 0  # Reset lower level sections
        #             current_sections[level_index] += 1  # Increment the corresponding section number
        #             section_title = '.'.join(str(num) for num in current_sections[:level_index + 1]) + " " + item.text
        #             chapter_structure.append({"title": section_title, "content": [], "tables": []})
        #         elif chapter_structure:  # Add paragraph content to the latest section
        #             chapter_structure[-1]["content"].append(item.text)

        #     elif isinstance(item, Table) and chapter_structure:
        #             chapter_structure[-1]["tables"].append(item)
        # print("chapter structure ok")
        # return chapter_structure
    
    def create_stream_from_doc(self) -> dict:
        """
        Create PDF streams from the document content.

        Returns:
            dict: A dictionary containing the PDF streams for each section.
        """
        chapters = self.parse_document()

        pdf_bytes_dict = {}

        for chapter in chapters:
            pdf_file_name = chapter["title"].replace(" ", "_").replace("\n", "") + ".pdf"
            
            try:
                pdf_buffer = io.BytesIO()
                doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
                styles = getSampleStyleSheet()
                flowables = [PDFParagraph(chapter["title"], styles['Heading1'])]

                for content in chapter["content"]:
                    flowables.append(PDFParagraph(content, styles['Normal']))

                for table in chapter["tables"]:
                    table_data = ""
                    for row in table.rows:
                        for cell in row.cells:
                            table_data += cell.text + '\t'
                        table_data += '\n'
                    flowables.append(PDFParagraph(table_data, styles['Normal']))

                doc.build(flowables)
                pdf_buffer.seek(0)
                pdf_bytes_dict[unquote(pdf_file_name)] = pdf_buffer.getvalue()

            except Exception as e:
                print(f"Error occurred while writing {pdf_file_name}: {str(e)}")
        return pdf_bytes_dict


    def draw_content_recursive(self, pdf, content, content_y, depth=0):
        """
        Recursively draw content on PDF canvas.

        Args:
            pdf: The PDF canvas.
            content: The content to draw.
            content_y: The y-coordinate for drawing content.
            depth: The depth of the content hierarchy.
        Returns:
            int: The updated y-coordinate after drawing content.
        """
        for item in content:
            if 'content' in item:
                content_y = self.draw_content_recursive(pdf, item['content'], content_y, depth + 1)
            else:
                pdf.drawString(20 + (depth * 10), content_y, item)
                content_y -= 12

        return content_y
    
    def create_dict(self) -> dict:
        """
        Create a dictionary representation of the document sections.

        Returns:
            dict: A dictionary containing the document sections and their content.
        """
        result_dict = {}
        sections = self.extract_sections()
        for section in sections:
            section_dict = {'title': section['title'], 'content': []}
            for subsection in section['content']:
                subsection_dict = {'title': subsection['title'], 'content': []}
                for subsubsection in subsection['content']:
                    subsubsection_dict = {'title': subsubsection['title'], 'content': []}
                    for subsubsubsection in subsubsection['content']:
                        subsubsubsection_dict = {'title': subsubsubsection['title'], 'content': []}
                        for paragraph in subsubsubsection['content']:
                            subsubsubsection_dict['content'].append(paragraph)
                        subsubsection_dict['content'].append(subsubsubsection_dict)
                        result_dict[subsubsubsection['title']] = subsubsubsection_dict
                    subsection_dict['content'].append(subsubsection_dict)
                    result_dict[subsubsection['title']] = subsubsection_dict
                section_dict['content'].append(subsection_dict)
                result_dict[subsection['title']] = subsection_dict

        return result_dict

    def extract_sections(self) -> list[dict]:
        """
        Extract sections from the document and organize them hierarchically.

        Returns:
            list[dict]: A list of dictionaries representing the sections and their content.
        """
        doc = self.pdf_to_docx(self.pdf_doc_stream)
        sections = []
        current_section = None
        current_subsection = None
        current_subsubsection = None
        current_subsubsubsection = None
        section_number = 0
        subsection_number = 1
        subsubsection_number = 1
        subsubsubsection_number = 1

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Check if paragraph has text
                #print(paragraph.text)
                if paragraph.style.name.startswith('Heading'):
                    heading_level = int(paragraph.style.name.split(' ')[-1])

                    if heading_level == 1:
                        section_number += 1
                        subsection_number = 0
                        subsubsection_number = 1
                        subsubsubsection_number = 1
                        current_section = {'title': f"{section_number}. {paragraph.text}", 'content': []}
                        sections.append(current_section)
                        current_subsection = None
                        current_subsubsection = None
                        current_subsubsubsection = None
                    elif heading_level == 2:
                        current_subsection = {'title': f"{section_number}.{subsection_number} {paragraph.text}",
                                            'content': []}
                        current_section['content'].append(current_subsection)
                        current_subsubsection = None
                        current_subsubsubsection = None
                        subsection_number += 1
                        subsubsection_number = 1
                        subsubsubsection_number = 1
                    elif heading_level == 3:
                        current_subsubsection = {'title': f"{section_number}.{subsection_number}.{subsubsection_number} "
                                                        f"{paragraph.text}", 'content': []}
                        current_subsection['content'].append(current_subsubsection)
                        current_subsubsubsection = None
                        subsubsection_number += 1
                        subsubsubsection_number = 1
                    elif heading_level == 4:
                        current_subsubsubsection = {'title': f"{section_number}.{subsection_number}.{subsubsection_number}"
                                                            f".{subsubsubsection_number} {paragraph.text}",
                                                    'content': []}
                        current_subsubsection['content'].append(current_subsubsubsection)
                        subsubsubsection_number += 1

                elif current_section is not None and current_subsection is not None \
                        and current_subsubsection is not None and current_subsubsubsection is not None:
                    current_subsubsubsection['content'].append(paragraph.text)
        return sections